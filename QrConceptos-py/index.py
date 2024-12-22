import pandas as pd
import qrcode
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Cm
import io
import os

# Ruta a la fuente TrueType. Asegúrate de que la ruta es correcta en tu sistema.
font_path = 'C:\\Windows\\Fonts\\arial.ttf'  # Cambia esta ruta según tu sistema

# Verificar si la fuente existe
if not os.path.exists(font_path):
    raise FileNotFoundError(f"No se encontró la fuente en la ruta especificada: {font_path}")

# Cargar el archivo CSV con punto y coma como delimitador
data = pd.read_csv('nombres.csv', delimiter=';')

# Imprimir los nombres de las columnas para verificar
print("Nombres de columnas después de cargar:", data.columns)

# Convertir todos los datos a tipo str para evitar errores de tipo
data = data.astype(str)

# Crear el documento de Word
doc = Document()

# Establecer márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)  # Margen superior de 2 cm
    section.bottom_margin = Cm(2)  # Margen inferior de 2 cm
    section.left_margin = Cm(0)  # Margen izquierdo de 0 cm
    section.right_margin = Cm(0)  # Margen derecho de 0 cm

# Añadir una tabla con 3 columnas
table = doc.add_table(rows=0, cols=3)  # Comenzar sin filas
table.autofit = False  # Desactivar el ajuste automático

# Configurar la fuente para el texto debajo de los QR
font_size = 18  # Aumentar tamaño de la fuente
font = ImageFont.truetype(font_path, font_size)

# Dividir los datos en grupos de 3 para que cada fila tenga tres elementos diferentes
rows_data = [data.iloc[i:i+3] for i in range(0, len(data), 3)]

for row_data in rows_data:
    # Añadir una nueva fila a la tabla
    row_cells = table.add_row().cells

    # Configurar la altura de las filas a 5 cm
    for cell in row_cells:
        cell.height = Cm(5)  # Altura de la celda 5 cm
        cell.vertical_alignment = 1  # Centrar verticalmente el contenido en la celda

    # Recorrer cada celda de la fila y añadir un QR para cada registro en `row_data`
    for i, (_, row) in enumerate(row_data.iterrows()):
        try:
            # Obtener los datos del registro actual
            qr_data = row['A']
            apellido = row['B']
            nombre = row['C']
            
            # Generar QR a partir del campo A
            qr = qrcode.make(qr_data)
            
            # Crear la imagen final con QR y texto
            text = f"{apellido}, {nombre}"
            qr_width, qr_height = qr.size
            
            # Calcular el tamaño del texto
            dummy_img = Image.new('RGB', (1, 1))
            draw_dummy = ImageDraw.Draw(dummy_img)
            text_bbox = draw_dummy.textbbox((0, 0), text, font=font)
            text_width = text_bbox[2] - text_bbox[0]
            
            # Definir el ancho de la imagen como el máximo entre el ancho del QR y el texto
            img_width = max(qr_width, text_width)
            img_height = qr_height + text_bbox[3]  # Sin margen entre QR y texto
           
            # Crear la imagen final
            text_img = Image.new('RGB', (img_width, img_height), 'white')
            
            # Pegar el QR en el centro
            qr_x = (img_width - qr_width) // 2
            text_img.paste(qr, (qr_x, 0))
            
            # Añadir el texto debajo del QR, centrado
            draw = ImageDraw.Draw(text_img)
            text_x = (img_width - text_width) // 2
            text_y = qr_height  # Colocar el texto justo debajo del QR
            draw.text((text_x, text_y), text, fill="black", font=font)
            
            # Guardar imagen como archivo temporal
            img_stream = io.BytesIO()
            text_img.save(img_stream, format="PNG")  # Usar PNG para mejor calidad
            img_stream.seek(0)

            # Añadir la imagen QR con texto en la celda actual
            paragraph = row_cells[i].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(img_stream, width=Inches(1.5))  # Tamaño del QR en el documento
            
            img_stream.close()

        except KeyError as e:
            print(f"Error: No se pudo acceder a la columna {e}")

# Guardar el documento
doc.save('resultado.docx')
print("Documento 'resultado.docx' generado exitosamente.")
